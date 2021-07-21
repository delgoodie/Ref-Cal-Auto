# !/usr/bin/python

# 99AA03-0221-9312


from os.path import exists as os_path_exists
from os import remove as os_remove
from os import listdir as os_listdir
from os import system as os_system
from os import rename as os_rename
from os import _exit as os__exit
from sys import exc_info as sys_exc_info
from re import search as re_search
from re import match as re_match
from shutil import copyfile as shutil_copyfile
import PySimpleGUI as sg
from docx import Document as docx_document
from docx.shared import Inches, Pt
import matplotlib as mpl
import matplotlib.pyplot as plt
from docx2pdf import convert as docx2pdf_convert
from math import floor, log10
from threading import Timer
from time import perf_counter
from datetime import date

# Globals

window = 0
params = 0
config = {}
internal_reqs = [
    # region Table II
    {"material": "Spectralon", "reflectance": 2, "geometry": "Puck", "tolerance": {600: (0, 2)}, "flatness": 4},
    {"material": "Spectralon", "reflectance": 2, "geometry": "Target", "tolerance": {600: 2}, "flatness": 4},
    {"material": "Spectralon", "reflectance": 5, "geometry": "Puck", "tolerance": {600: 1}, "flatness": 4},
    {"material": "Spectralon", "reflectance": 5, "geometry": "Target", "tolerance": {600: 2}, "flatness": 4},
    {"material": "Spectralon", "reflectance": (6, 19), "geometry": "Puck", "tolerance": {600: 1}, "flatness": 5},
    {"material": "Spectralon", "reflectance": (6, 19), "geometry": "Target", "tolerance": {600: 3}, "flatness": 5},
    {"material": "Spectralon", "reflectance": (20, 95), "geometry": "Puck", "tolerance": {600: 2}, "flatness": 5},
    {"material": "Spectralon", "reflectance": (20, 95), "geometry": "Target", "tolerance": {600: 5}, "flatness": 5},
    # endregion
    # region Table III
    {
        "material": "Spectralon",
        "reflectance": 99,
        "geometry": "Target",
        "range": {
            250: (0.9, 0.995),
            300: (0.925, 0.995),
            350: (0.975, 0.995),
            (400, 700, 50): (0.985, 0.995),
            (750, 1600, 50): (0.975, 0.995),
            (1650, 2350, 50): (0.915, 0.995),
        },
    },
    # endregion
    # region Permaflect Table
    {"material": "Permaflect", "reflectance": 5, "geometry": "Target", "tolerance": {905: 1.25}},
    {"material": "Permaflect", "reflectance": 10, "geometry": "Target", "tolerance": {905: 1.25}},
    {"material": "Permaflect", "reflectance": 18, "geometry": "Target", "tolerance": {905: 1.25}},
    {"material": "Permaflect", "reflectance": 50, "geometry": "Target", "tolerance": {905: 1.75}},
    {"material": "Permaflect", "reflectance": 80, "geometry": "Target", "tolerance": {905: 1.25}},
    {"material": "Permaflect", "reflectance": 94, "geometry": "Target", "tolerance": {905: (0, 3)}}
    # endregion
]

# region HELPERS


class Date:
    def __init__(self, month: int, day: int, year: int):
        self.month = month
        self.day = day
        self.year = year

    def valid(self):
        return self.month >= 1 and self.month <= 12 and self.day > 0 and self.day <= 31 and self.year > 2000 and self.year < 3000

    def __eq__(self, other):
        return type(other) is Date and self.month == other.month and self.day == other.day and self.year == other.year

    def __str__(self):
        return f"{self.month}/{self.day}/{self.year}"


class Parameters:
    def __init__(
        self,
        root_path: str = None,
        geometry: str = None,
        size: str = None,
        material: str = None,
        serial_number: str = None,
        reflectance: int = None,
        nvlap: bool = False,
        requirements: dict = None,
        instrument: str = None,
        date: Date = None,
        stray_light_path: str = None,
    ):
        if len(root_path) > 0 and not root_path[-1] in ["\\", "/"]:
            root_path += "\\"
        if type(reflectance) is str:
            search = re_search("(\\d+)\\%", reflectance)
            if search:
                reflectance = int(search.group(1))
        if type(date) is str:
            date = DateFromString(date)
        if len(stray_light_path) > 0 and not stray_light_path[-1] in ["\\", "/"]:
            stray_light_path += "\\"

        self.root_path = root_path
        self.geometry = geometry
        self.size = size
        self.material = material
        self.serial_number = serial_number
        self.reflectance = reflectance
        self.nvlap = nvlap
        self.requirements = requirements
        self.instrument = instrument
        self.date = date
        self.stray_light_path = stray_light_path

        # derived property
        self.model = f"{'SRT' if self.geometry == 'Target' else 'SRS'}-{self.reflectance}-{self.size}"
        self.docx_name = ""
        if self.reflectance == 99:
            if self.nvlap:
                self.docx_name = "DM-01400-001Rev13 99 cal cert.docx"
            else:
                self.docx_name = "User Data\\DM-01400-009Rev04 99 cal cert non NVLAP.docx"
        else:
            if self.nvlap:
                self.docx_name = "User Data\\DM-01400-001Rev13 Gray cal cert.docx"
            else:
                self.docx_name = "User Data\\DM-01400-009Rev04 Gray cal cert non NVLAP.docx"

    def isValid(self):
        if not (self.root_path and type(self.root_path) is str and os_path_exists(self.root_path)):
            return "Invalid Root Path"
        if not (self.material and type(self.material) is str and self.material in ["Spectralon", "Permaflect"]):
            return "Invalid Material"
        if not (self.geometry and type(self.geometry) is str and self.geometry in ["Target", "Puck"]):
            return "Invalid Type"
        if not (self.size and type(self.size) is str):
            return "Invalid [Puck or Target] Size"
        if not (self.serial_number and type(self.serial_number) is str and len(self.serial_number) > 4):
            return "Invalid Serial Number"
        if not (self.reflectance and type(self.reflectance) is int and self.reflectance > 0 and self.reflectance < 100):
            return "Invalid Reflectance"
        if not (self.instrument and self.instrument in ["A", "B", "C"]):
            return "Invalid Instrument"
        if not (self.date and type(self.date) is Date and self.date.valid()):
            return "Invalid Date"
        if not (self.stray_light_path and type(self.stray_light_path) is str and os_path_exists(self.stray_light_path)):
            return "Invalid Stray Light Path"
        return True


class CSV:
    def __init__(self, _path: str):
        self.path = _path
        file = open(self.path, "r")
        self._data = []
        for l in file.read().split("\n"):
            self._data.append(l.split(","))
        file.close()
        self.modified = False

    def _ParseTablePosition(self, position) -> tuple[int, int]:
        if type(position) is str:
            if ":" in position:
                c, r = position.split(":", 1)
            else:
                match = re_match("([a-zA-Z]+)(\\d+)", position)
                if len(match.regs) > 1:
                    c = match.group(1)
                    r = match.group(2)
            if c and r:
                position = (ord(c.upper()) - 65, int(r) - 1)
            else:
                raise ValueError(f"CSV._ParseTablePosition: c:{c}, r:{r} is invalid position")
        elif type(position) is tuple:
            c, r = position
            if type(c) is int or type(c) is float or re_match("\\d+", c):
                c = int(c)
            elif re_match("[A-Za-z]+", c):
                c = c.upper()
                col = 0
                i = 1
                for char in c:
                    col += (ord(char) - 65) * i
                    i *= 26
                c = col
            r = int(r)
            position = (c, r - 1)
        else:
            raise TypeError(f"CSV._ParseTablePosition: position: {type(position)} is not <class 'str'> or <class 'tuple'>")

        if type(position) is tuple and type(position[0]) is int and type(position[1]) is int:
            return position
        else:
            raise ValueError(f"CSV._ParseTablePosition: c:{c}, r:{r} is invalid position")

    def Read(self, position="*"):
        """
        Read from csv table cell(s)

        position -- cell position: '*', (0, 34) | ('A', 34) | ('A', '34') | 'A34' | 'A:34'
        if position is '*' then a 2D array of all cols & rows will be returned

        return -- data at position
        """
        if position == "*":
            return self._data
        else:
            parsed_pos = self._ParseTablePosition(position)
            if parsed_pos[1] < len(self._data) and parsed_pos[0] < len(self._data[parsed_pos[0]]):
                return self._data[parsed_pos[1]][parsed_pos[0]]
            else:
                raise IndexError(f"easyio.CSV.Read: position:{parsed_pos} not in cols: {len(self._data[0])}, rows: {len(self._data)}")

    def Write(self, position, data: str = ""):
        self.modified = True
        """
        Write string to single csv table cell

        position -- cell position: (0, 34) | ('A', 34) | ('A', '34') | 'A34' | 'A:34'
        data -- string of new cell contents: '34.213'

        return -- the overwriten data: 'data before Write()'
        """
        parsed_pos = self._ParseTablePosition(position)

        pre = self._data[parsed_pos[1]][parsed_pos[0]]
        self._data[parsed_pos[1]][parsed_pos[0]] = data
        return pre

    def __del__(self):
        if self.modified:
            string = ""
            for c in self._data:
                for r in c:
                    string += r + ","
                string += "\n"

            file = open(self.path, "w")
            file.write(string)
            file.close()


class DOCX:
    def __init__(self, _path):
        self.path = _path
        self.doc = docx_document(_path)

    def _docxOccurences(self, variable: str) -> None:
        occurences = []
        for p in self.doc.paragraphs:
            if re_search(f"<{variable}>", p.text):
                occurences.append(p)
        for t in self.doc.tables:
            for c in t._cells:
                if re_search(f"<{variable}>", c.text):
                    occurences.append(c.paragraphs[0])
        for s in self.doc.sections:
            for p in s.header.paragraphs:
                if re_search(f"<{variable}>", p.text):
                    occurences.append(p)
            for t in s.header.tables:
                for c in t._cells:
                    if re_search(f"<{variable}>", c.text):
                        occurences.append(c.paragraphs[0])
        return occurences

    def ReplaceText(self, variable: str, value: str):
        for e in self._docxOccurences(variable):
            text = e.text.replace(f"<{variable}>", value)
            e.text = ""
            run = e.add_run()
            run.text = text
            run.font.name = "Times New Roman"
            run.font.size = Pt(8)

    def ReplacePicture(self, variable: str, path: str, size):
        for e in self._docxOccurences(variable):
            e.text = ""
            e.add_run().add_picture(path, width=Inches(size[0]), height=Inches(size[1]))

    def Save(self, path=None):
        if path:
            self.doc.save(path)
        else:
            self.doc.save(self.path)


def DateFromString(string: str) -> Date:
    """
    Retrieves date from string containing a date in the following formats:

    MM/DD/YY(YY)

    MM-DD-YY(YY)

    NOTE: Shorthand and Full years will be converted to full years

    i.e. 21 => 2021, 1985 => 1985

    @return tuple containing (MM, DD, YYYY)
    """
    search = re_search("(\\d+)[\\/\\-](\\d+)[\\/\\-](\\d+)", string)
    if not search or len(search.regs) != 4:
        return -1
    else:
        month = int(search.group(1))
        day = int(search.group(2))
        year = int(search.group(3))
        # check if year is in shorthand, '21, or longhand, 2021 -> convert to longhand
        if log10(year) <= 3:
            year += 2000
    return Date(month, day, year)


def LeftPad(string: str, length: int) -> str:
    """
    Left Pad a string with ' ' (spaces) so that it takes on the length of @param length
    """
    while len(string) < length:
        string += " "
    return string


def PreTabCount(string: str) -> int:
    i = 0
    count = 0
    while string[i] in [" ", "\t"]:
        if string[i : i + 4] == "    " or string[i : i + 4] == "   \t":
            i += 4
        elif string[i : i + 3] == "  \t":
            i + 3
        elif string[i : i + 2] == " \t":
            i += 2
        elif string[i] == "\t":
            i += 1
        else:
            raise Exception("main.PreTabCount(string: str)->int Invalid Tabs and Spaces")
        count += 1
    return count


def ParseTabTree(lines: list[str]) -> dict:
    root = {}
    i = 0
    while i < len(lines):
        if lines[i] == "\n" or lines[i][0] == "#":
            i += 1
        elif "=" in lines[i]:
            root[lines[i].split("=")[0].strip()] = lines[i].split("=")[1].strip()
            i += 1
        else:
            num_tabs = PreTabCount(lines[i])
            j = i + 1
            while j < len(lines) and (PreTabCount(lines[j]) > num_tabs or lines[j] == "\n"):
                j += 1
            root[lines[i].strip()] = ParseTabTree(lines[i + 1 : j])
            i = j
    return root


def debug(func):
    """
    Debug Wrapper to catch errors, display timestamp, and log status

    use:

    @debug

    def MyFunc
    """
    global window

    def wrap(*args, **kwargs):
        start = perf_counter()
        print(LeftPad(func.__name__, 40), end="")
        window["Log"].update(func.__name__)
        result = False
        try:
            result = func(*args, **kwargs)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys_exc_info()
            print("\n\n[Exception Raised]")
            print("Send email to:    wdelgiudice@labsphere.com    with the following information:")
            print(
                "* error message (seen below) \n* all form data (model, sn, date, requirements, etc)\n* anything different about this specific scan"
            )
            print("Error Message:")
            print("TYPE: ", exc_type)
            print("LINE: ", exc_tb.tb_lineno)
            print("ERROR: ", e)
            window["Log"].update("ERROR: GO TO CONSOLE")
            return False
        end = perf_counter()
        print(f"√  {round(1000 * (end - start))} ms")
        window["Log"].update(f"{func.__name__}   √")
        return result

    return wrap


@debug
def GetStrayLightPaths(date: Date) -> list[str]:
    global config
    """
    Returns array of relative paths from STRAY_LIGHT_DIR to stray light folders written at dates that match the @param date

    @param date - must be tuple in (DD:int, MM:int, YYYY:int) format

    @return - List[str] of relative paths from STRAY_LIGHT_DIR
    """
    result = []
    for dir in os_listdir(config["stray light directory"]):
        dirDate = DateFromString(dir)
        if dirDate != -1 and dirDate == date:
            result.append(dir)
    return result


# endregion

# region EXECUTION STEPS


@debug
def GetDocxTemplate() -> DOCX:
    global params
    return DOCX(f"User Data\\{params.docx_name}")


@debug
def Get_rr() -> list[float]:
    """
    Reads Rr data from \\User Data\\rr.txt
    """
    return [float(l) for l in open("User Data\\rr.txt").readlines()]


@debug
def CorrectData(raw: DOCX, strayLight: DOCX, Rr: list[float]) -> list[float]:
    """
    Calculates Corrected data from raw data, stray light data, and Rr
    """
    # MS% in column B
    Ms = {}
    i = 2
    while len(raw.Read(("A", i))) > 0:
        Ms[int(float(raw.Read(("A", i))))] = float(raw.Read(("B", i)))
        i += 1

    Mh = {}
    i = 2
    while len(strayLight.Read(("A", i))) > 0:
        s_l = float(strayLight.Read(("B", i)))
        if s_l < 0 or s_l > 1:
            print(f"[WARNING] Check Stray Light Scan for data outside of range 0 - 1 (around {i + 245}nm)")
            s_l = max(0, min(1, s_l))
        Mh[int(float(strayLight.Read(("A", i))))] = s_l
        i += 1

    c_d = {}
    for w in Mh:
        if w >= 250 and w <= 2500:
            c_d[w] = round((Ms[w] * 0.01 - Mh[w] * 0.01) * (1 / (1 - Mh[w] * 0.01)) * Rr[w - 250], 4)
            if c_d[w] < 0:
                raise Exception(f"negative corrected data: {c_d[w]} @ {w}")

    return c_d


@debug
def TestRequirements(corrected_data: dict) -> str:
    """
    Tests Corrected Data against Requirements

    @return - 0 if data passed tests, str if data failed for error msg
    """
    global params, internal_reqs

    # Test Internal Requirements

    for i_r in internal_reqs:
        if params.material == i_r["material"] and params.reflectance == i_r["reflectance"] and params.geometry == i_r["geometry"]:
            if "range" in i_r:
                for r in i_r["range"]:
                    if type(r) is int:
                        if corrected_data[r] < i_r["range"][r][0] or corrected_data[r] > i_r["range"][r][1]:
                            return f"Corrected Data did not pass requirements (internal) ({corrected_data[r]} @ {r}nm did not meet {i_r['range'][r][0]} < R < {i_r['range'][r][1]} @ {r}nm)"
                    else:
                        for w in range(r[0], r[1] + r[2], r[2]):
                            if corrected_data[w] < i_r["range"][r][0] or corrected_data[w] > i_r["range"][r][1]:
                                return f"Corrected Data did not pass requirements (internal) ({corrected_data[w]} @ {w}nm did not meet {i_r['range'][r][0]} < R < {i_r['range'][r][1]} on [{r[0]}nm - {r[1]}nm))"
            if "tolerance" in i_r:
                target = 0
                if params.material == "Spectralon":
                    target = params.reflectance
                else:
                    target = sum(corrected_data[v] for v in corrected_data) / len(corrected_data) * 100
                for t in i_r["tolerance"]:
                    r = i_r["tolerance"][t]
                    if type(t) is int:
                        if type(r) is int or type(r) is float:
                            if corrected_data[t] < (target - r) * 0.01 or corrected_data[t] > (target + r) * 0.01:
                                return f"Corrected Data did not pass requirements (internal) ({corrected_data[t]} @ {t}nm did not meet R within +/-{r}% of {target} @ {t}nm)"
                        else:
                            if corrected_data[t] < (target - r[0]) * 0.01 or corrected_data[t] > (target + r[1]) * 0.01:
                                return f"Corrected Data did not pass requirements (internal) ({corrected_data[t]} @ {t}nm did not meet R within -{r[0]}% to +{r[1]}% of {target} @ {t}nm)"

                    else:
                        for w in range(t[0], t[1] + t[2], t[2]):
                            if type(r) is int or type(r) is float:
                                if corrected_data[w] < (target - r) * 0.01 or corrected_data[t] > (target + r) * 0.01:
                                    return f"Corrected Data did not pass requirements (internal) ({corrected_data[w]} @ {w}nm did not meet R within +/-{r}% of {target} on [{r[0]}nm - {r[1]}nm))"
                            else:
                                if corrected_data[t] < (target - r[0]) * 0.01 or corrected_data[t] > (target + r[1]) * 0.01:
                                    return f"Corrected Data did not pass requirements (internal) ({corrected_data[w]} @ {w}nm did not meet R within -{r[0]}% to +{r[1]}% of {target} on [{r[0]}nm - {r[1]}nm))"
            if "flatness" in i_r:
                flat_arr = [corrected_data[v] for v in range(350, 760)]
                flatness = max(flat_arr) - min(flat_arr)
                if flatness > i_r["flatness"]:
                    return f"Corrected Data did not pass requirements (internal) (flatness: {flatness} was not < .02)"

    # Test Additional Requirements

    for req in params.requirements:
        if type(req) is int:
            if corrected_data[req] < params.requirements[req][0] or corrected_data[req] > params.requirements[req][1]:
                return f"Corrected Data did not pass requirements (additional) ({corrected_data[req]} @ {req} did not meet {params.requirements[req][0]} <= reflectance <= {params.requirements[req][1]} @ {req})"
        elif type(req) is tuple:
            for w in range(req[0], req[1]):
                if w in corrected_data and (corrected_data[w] < params.requirements[req][0] or corrected_data[w] > params.requirements[req][1]):
                    return f"Corrected Data did not pass requirements (additional) ({corrected_data[w]} @ {w} did not meet {params.requirements[req][0]} <= reflectance <= {params.requirements[req][1]} @ [{req[0]}, {req[1]}])"
        elif req == "flatness":
            flat_arr = [corrected_data[v] for v in range(350, 760)]
            flatness = max(flat_arr) - min(flat_arr)
            if flatness > params.requirements[req]:
                return f"Corrected Data did not pass requirements (additional) (flatness: {flatness} was not < {params.requirements[req]})"


@debug
def RenameRootFolder(success: bool) -> None:
    """
    Renames root folder, either sn if success or sn-FAIL

     removing last four sn digits and appending FAIL

    If a FAIL folder already exists, append a number starting at 2 and increasing by one:

    -FAIL-2, -FAIL-3, -FAIL-{n}
    """
    global params

    path = params.root_path[0 : params.root_path.rindex("/") + 1] + params.serial_number
    if not success:
        if not os_path_exists(f"{path[0:-5]}FAIL"):
            path = f"{path[0:-5]}FAIL"
        else:
            i = 2
            while os_path_exists(f"{path[0:-5]}FAIL-{i}"):
                i += 1
            path = f"{params.root_path[0:-5]}FAIL-{i}"

    os_rename(params.root_path, path)
    params.root_path = path + "\\"


@debug
def SaveStrayLight(src_path):
    global params
    shutil_copyfile(src_path, f"{params.root_path}\\StrayLightScan.csv")


@debug
def SaveTextFile(corrected_data: dict) -> None:
    """
    Saves corrected data as text file under (last four of sn)-(model name).txt
    """
    # print(f"Generating {params['serial number'][-4:len(params['serial number'])]}-{params['model']}.txt file    ")
    txt = open(f"{params.root_path}{params.serial_number[-4:len(params.serial_number)]}-{params.model}.txt", "w")
    stringdata = [f"{w}\t{corrected_data[w]}\n" for w in range(250, 2501)]
    stringdata.insert(0, f"{params.serial_number}\nThis data is for reference only\n")
    txt.write("".join(stringdata))


@debug
def WriteWordMeta(doc: DOCX) -> None:
    """
    Writes metadata to word docx template

    * sn
    * date
    * model
    * instrument (A, B, C)
    """
    doc.ReplaceText("sn", params.serial_number)
    doc.ReplaceText("DATE", date.today().strftime("%m/%d/%Y"))
    doc.ReplaceText("model", params.model)
    doc.ReplaceText("isA", "X" if params.instrument == "A" else "")
    doc.ReplaceText("isB", "X" if params.instrument == "B" else "")
    doc.ReplaceText("isC", "X" if params.instrument == "C" else "")


@debug
def WriteWordData(doc: DOCX, corrected_data: dict) -> None:
    """
    Writes corrected reflectance data to word docx cert in in table

    Rounds reflectance to nearest value in uncertainty chart and uses values for sig fig rounding
    """
    uncertainty_table = doc.doc.tables[2]._cells
    uncertainty = {}
    offset = -1
    uncertainty_ref = 0
    UNCERTAINTY_COLS = 9

    # Find nearest reflectance in Uncertainty Table
    for i in range(1, UNCERTAINTY_COLS):
        col_ref = int(re_search("(\\d+)%", uncertainty_table[i].text).group(1))
        if params.reflectance == col_ref or not uncertainty_ref or abs(params.reflectance - col_ref) < abs(params.reflectance - uncertainty_ref):
            uncertainty_ref = col_ref
            offset = i

    # Find sigfigs for each wavelength from Uncertainty Table
    for i in range(UNCERTAINTY_COLS + offset, len(uncertainty_table), UNCERTAINTY_COLS):
        w = int(uncertainty_table[i - offset].text)
        u = 0
        non_zero = False
        uncert = uncertainty_table[i].text
        for c in uncertainty_table[i].text:
            if c in "123456789":
                u += 1
                non_zero = True
            elif c == "0" and non_zero:
                u += 1
        uncertainty[w] = u

    # Apply sigfigs to corrected_data and write result to Docx
    for w in range(250, 2510, 50):
        if w in corrected_data:
            v = corrected_data[w]
            u = uncertainty[w]
            if v and u:
                # rounding to sig figs based off uncertainty table
                v = round(v, u - int(floor(log10(abs(v)))))
                v = str(v)
                while len(v) < u + 3:
                    v += "0"
                doc.ReplaceText(f"w{str(int(w / 10))}", v)
            else:
                print(f"No corrected_data at wavelength: {w}")


@debug
def WriteWordGraph(doc: DOCX, corrected_data: dict) -> None:
    """
    Writes graph to word docx cert

    * Creates graph with matlib from corrected data
    * Saves graph to temp.png
    * Writes image to docx cert
    * Deletes temp.png
    """
    mpl.use("Agg")
    plt_x = [w for w in range(250, 2500, 5)]
    plt_y = []
    for w in plt_x:
        avg = 0
        for i in range(w, w + 5):
            avg += corrected_data[i]
        avg /= 5
        plt_y.append(avg)
    plt.plot(plt_x, plt_y, color="black")
    # plt.title("Graph I: 8°/Hemispherical Spectral Reflectance")
    plt.ylabel("Reflectance Factor")
    plt.xlabel("Wavelength (nm)")
    plt.xticks([i for i in range(250, 2501, 250)])
    # plt.axis([250, 2500, 0, ceil(max(plt_y) * 10) * 0.1])
    plt.savefig("temp.png", format="png")
    doc.ReplacePicture("graph", "temp.png", (7, 5.5))
    os_remove("temp.png")


@debug
def SaveWord(doc: DOCX) -> None:
    """
    Saves word doc cert

    This function exists simply to wrap DOCX.Save() method so the debug wrapper can be used
    """
    doc.Save(params.root_path + doc.path[doc.path.rindex("\\") + 1 : len(doc.path)])


@debug
def SavePdf() -> None:
    """
    Copies word doc cert as pdf

    This function exists simply to wrap docx2pdf_covert() method so the debug wrapper can be used
    """
    print("")
    docx2pdf_convert(
        f"{params.root_path}{params.docx_name}",
        f"{params.root_path}{params.serial_number}.pdf",
    )


@debug
def CopyToUsb() -> None:
    global config
    """
    Copies raw data txt file and final cert pdf file to USB path specified in User Data\\config.txt
    """
    if os_path_exists(config["usb path"]):
        txtName = f"{params.serial_number[-4:len(params.serial_number)]}-{params.model}.txt"
        shutil_copyfile(f"{params.root_path}{txtName}", f"{config['usb path']}{txtName}")
        shutil_copyfile(f"{params.root_path}{params.serial_number}.pdf", f"{config['usb path']}{params.serial_number}.pdf")


# endregion

# region EXECUTION


def Execute() -> bool:
    """
    Execute function sequentially calls steps and exits if any step fails:

    Execute is done in many steps so that each step can be wrapped with @debug for easy debugging if any point fails

    * Get Rr
    * Corrected Data
    * Test Client Requirements
    * Save Text File
    * Write Word Meta
    * Write Word Data
    * Write Word Graph
    * Save Word
    * Save Pdf
    * Copy To Usb
    """
    global params
    global window

    raw = CSV(f"{params.root_path}Equation1.Sample.Cycle1.Equation1.csv")
    strayLight = CSV(f"{params.stray_light_path}Equation1.Sample.Cycle1.Equation1.csv")

    doc = GetDocxTemplate()

    rr = Get_rr()
    if not rr:
        return False

    corrected_data = CorrectData(raw, strayLight, rr)
    if not corrected_data:
        return False

    msg = TestRequirements(corrected_data)
    if msg:
        RenameRootFolder(False)
        window["Log"].update(msg)
        print(msg)
        return False
    else:
        RenameRootFolder(True)

    SaveStrayLight(strayLight.path)

    SaveTextFile(corrected_data)

    WriteWordMeta(doc)

    WriteWordData(doc, corrected_data)

    WriteWordGraph(doc, corrected_data)

    SaveWord(doc)

    SavePdf()

    CopyToUsb()

    return True


def AsyncExecute() -> None:
    """
    This function exists so the Execute Function can occur on a separate thread, while the GUI continues to receive updates
    """
    global window
    global params
    global config

    result = Execute()
    if result:
        window["Log"].update("Finished: SUCCESS")
        os_system(f'start "" "{params.root_path}"')
        if os_path_exists(config["usb path"]):
            os_system(f'start "" "{config["usb path"]}"')
        window.close()
        os__exit(0)


# endregion

# region EVENTS


def GeometryEvent(values):
    global window

    isTarget = values["Geometry"] == "Target"
    window["Size Name"].update("Target Size" if isTarget else "Puck Diameter")
    sizeDropdown = ["020", "050", "100", "120", "180", "240"] if isTarget else ["010", "020"]
    window["Size"].update(values=sizeDropdown)


def DateEvent(values):
    global window

    date = DateFromString(values["Date"])
    if date != -1:
        slps = GetStrayLightPaths(date)
        if len(slps) == 0:
            window["Stray Light Dropdown"].update(values=[])
        else:
            window["Stray Light Dropdown"].update(values=slps)


def StrayLightDropdownEvent(values):
    global window

    window["Stray Light Path"].update(f"{config['stray light directory']}{values['Stray Light Dropdown']}")
    # Suggest Instrument from Stray Light Path
    regex = re_search("[- ](A|B|C)( |-|$)", values["Stray Light Dropdown"])
    if regex:
        window["Instrument"].update(regex.groups(1)[0])


def ExecuteEvent(values):
    global window
    global params

    params = Parameters(
        values["Browse"],
        values["Geometry"],
        values["Size"],
        values["Material"],
        values["Serial Number"],
        values["Reflectance"],
        values["Nvlap"],
        config["requirements"][values["Requirements"]] if values["Requirements"] in config["requirements"] else {},
        values["Instrument"],
        values["Date"],
        values["Stray Light Path"],
    )
    error = params.isValid()
    if not type(error) is str:
        window["Log"].update("Executing...")
        Timer(1.0, AsyncExecute).start()
    else:
        window["Log"].update(error)
        print("Execution Conditions not met:     ", error, "\n")


# endregion


def setup() -> None:
    global window
    global params
    global config
    # checks for local folder User Data
    if not os_path_exists("User Data\\rr.txt"):
        raise Exception("User Data\\rr.txt Does Not Exist")
    if not os_path_exists("User Data\\config.txt"):
        raise Exception("User Data\\config.txt Does Not Exist")
    if not os_path_exists("User Data\\DM-01400-001Rev13 99 cal cert.docx"):
        raise Exception("User Data\\DM-01400-001Rev13 99 cal cert.docx Does Not Exist")
    if not os_path_exists("User Data\\DM-01400-001Rev13 Gray cal cert.docx"):
        raise Exception("User Data\\DM-01400-001Rev13 Gray cal cert.docx Does Not Exist")
    if not os_path_exists("User Data\\DM-01400-009Rev04 99 cal cert non NVLAP.docx"):
        raise Exception("User Data\\DM-01400-009Rev04 99 cal cert non NVLAP.docx Does Not Exist")
    if not os_path_exists("User Data\\DM-01400-009Rev04 Gray cal cert non NVLAP.docx"):
        raise Exception("User Data\\DM-01400-009Rev04 Gray cal cert non NVLAP.docx Does Not Exist")

    # Reads config data for constant paths

    parsed_config = ParseTabTree(open("User Data\\config.txt").readlines())

    config["target directory"] = parsed_config["REF_CAL_PATH"]
    config["stray light directory"] = parsed_config["STRAY_LIGHT_PATH"]
    config["usb path"] = parsed_config["USB_PATH"]

    # REQUIREMENTS
    int_req = {}
    for req in parsed_config["REQUIREMENTS"]:
        int_req[req] = {}
        for spec in parsed_config["REQUIREMENTS"][req]:
            if re_match("^\\d+$", spec):
                key = int(spec)
            elif re_match("\\d+\\-\\d+", spec):
                key = (int((spec.split("-")[0].strip())), int((spec.split("-")[1].strip())))
            else:
                key = spec

            raw = parsed_config["REQUIREMENTS"][req][spec]
            if re_match("\\d*\\.\\d+\\-\\d*\\.\\d+", raw):
                val = (float((raw.split("-")[0].strip())), float((raw.split("-")[1].strip())))
            elif re_match("\\d*\\.\\d+", raw):
                val = float(raw)
            else:
                val = raw
                print(raw)

            int_req[req][key] = val

    config["requirements"] = int_req

    layout = [
        # 0
        [sg.Text("Root Folder", font="30px")],
        # 1
        [sg.FolderBrowse(), sg.Input("(must have Equation1.Sample.Cycle1.Equation1.csv)", size=(82, 1), readonly=True)],
        # 2
        [sg.Text("")],
        # 3
        [sg.Text("Info", font="30px")],
        # 4
        [
            sg.Text("Geometry"),
            sg.DropDown(["Target", "Puck"], "Select", key="Geometry", size=(9, 1), readonly=True, enable_events=True),
            sg.Text("Material"),
            sg.DropDown(["Spectralon", "Permaflect"], "Select", key="Material", size=(13, 1), readonly=True, enable_events=True),
            sg.Text("Target Size", size=(11, 0), key="Size Name", pad=(0, 0)),
            sg.DropDown(["020", "050", "100", "120", "180", "240"], "", key="Size", size=(6, 1), pad=(0, 0), readonly=False),
            sg.Text("Reflectance"),
            sg.DropDown(
                ["2%", "5%", "10%", "18%", "20%", "40%", "50%", "60%", "75%", "80%", "99%"], "99%", size=(7, 1), key="Reflectance", readonly=False
            ),
        ],
        # 5
        [
            sg.Text("Serial Number", size=(10, 1)),
            sg.Input("", size=(20, 1), key="Serial Number"),
            sg.Checkbox("NVLAP", key="Nvlap", default=False),
            sg.Text("Requirements", pad=((0, 0), 0)),
            sg.DropDown([c for c in config["requirements"]], "No Additional Requirements", size=(30, 1), key="Requirements", readonly=True),
        ],
        # 6
        [sg.Text("")],
        # 7
        [sg.Text("Stray Light Scan", font="30px")],
        # 8
        [
            sg.Text("Instrument"),
            sg.DropDown(["A", "B", "C"], "A", size=(5, 1), key="Instrument", readonly=True),
            sg.CalendarButton(
                "Select Date",
                target=(8, 3),
                format="%m/%d/%Y",
                enable_events=True,
            ),
            sg.Input(key="Date", size=(10, 1), enable_events=True, readonly=True),
            sg.DropDown([], "Select Scan", size=(45, 1), enable_events=True, key="Stray Light Dropdown", readonly=True),
        ],
        # 9
        [
            sg.FolderBrowse(button_text="Manual Browse", target=(9, 1)),
            sg.Input("Stray Light Path (not selected)", key="Stray Light Path", size=(75, 1), readonly=True),
        ],
        # 10
        [sg.Text("")],
        # 11
        [sg.Button("Execute", size=(20, 1), font="30px", pad=(210, 0))],
        # 12
        [sg.Text("Log", font="30px")],
        # 13
        [sg.Input(key="Log", size=(91, 5), readonly=True, enable_events=True)],
    ]

    window = sg.Window(title="Ref Cal Auto", layout=layout, margins=(0, 20))


def main() -> None:
    """
    Main function:

    * Handles Events

    * Executes tool
    """
    global window
    global params
    global config

    while True:
        event, values = window.read()

        if event == "Geometry":
            GeometryEvent(values)
        elif event == "Material":
            0
        elif event == "Date":
            DateEvent(values)
        elif event == "Stray Light Dropdown":
            StrayLightDropdownEvent(values)
        elif event == "Execute":
            ExecuteEvent(values)
        elif event == sg.WIN_CLOSED:
            break
    window.close()


# ENTRY POINT
setup()
main()
